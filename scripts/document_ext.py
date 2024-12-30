import json
import os
from docx import Document
import pdfplumber
from PIL import Image
import pytesseract
import pandas as pd
from utils.utility import doc_to_pdf
from docling.document_converter import DocumentConverter
from pdf2image import convert_from_path
import re

converter = DocumentConverter()

cwd = os.getcwd()

class Data_extractor:
    def __init__(self):
        self.data = {}
    
    def format_tables(self, tb_lst):
        seg_ind = []
        for ind_m, val1 in enumerate(tb_lst):
            for xy in val1:
                if len(xy)>1:
                    check_str = xy[0]+xy[1]
                elif len(xy)==1:
                    check_str = xy[0]

                if (check_str in "1. PROJECT DATA") or (check_str in "GENERAL INFORMATION"):
                    seg_ind.append([ind_m])
                elif (check_str in "2. MARKETING INFORMATION") or (check_str in '2.1 Describe proposed products and their applications'):
                    seg_ind[0].append(ind_m-1)
                    seg_ind.append([ind_m])
                elif (check_str in "3. TECHNICAL INFORMATION") or (check_str in '3.1 Elaborate in details the manufacturing process descriptions'):
                    seg_ind[1].append(ind_m-1)
                    seg_ind.append([ind_m])
                elif (check_str in "4. FINANCIAL INFORMATION") or (check_str in '4.1 SOURCES OF FINANCE'):
                    seg_ind[2].append(ind_m-1)
                    seg_ind.append([ind_m])
                    seg_ind[3].append(len(tb_lst) - 1)
        
        try:
            tables_segments = {}
            tables_segments["1. PROJECT DATA"] = tb_lst[seg_ind[0][0]:seg_ind[0][1]]
            tables_segments["2. MARKETING INFORMATION"] = tb_lst[seg_ind[1][0]:seg_ind[1][1]]
            tables_segments["3. TECHNICAL INFORMATION"] = tb_lst[seg_ind[2][0]:seg_ind[2][1]]
            tables_segments["4. FINANCIAL INFORMATION"] = tb_lst[seg_ind[3][0]:seg_ind[3][1]]
            return tb_lst
        except IndexError:
            return []

    def loan_application_parser(self, file_path):
        file, extension = os.path.splitext(file_path)
        extension = extension.lower()
        tb_lst = []
        if extension in ['.doc', '.docx']:
            print("test-1")
            pdf_path = file + ".pdf"
            rs = doc_to_pdf(input_docx=file_path, output_pdf=pdf_path)
            
            if rs==True:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        tables = page.extract_tables()
                        for table in tables:
                    #         # Clean up table data
                            cleaned_table = [[cell.strip() if cell else "" for cell in row] for row in table]
                            tb_lst.append(cleaned_table)

            elif rs=="timeout":
                doc_path = file_path
                doc = Document(doc_path)
                # Iterate over all tables in the document
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    tb_lst.append(table_data)

        elif extension.lower() in [".pdf"]:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                #         # Clean up table data
                        cleaned_table = [[cell.strip() if cell else "" for cell in row] for row in table]
                        tb_lst.append(cleaned_table)

        fmt_tbls = self.format_tables(tb_lst)
        
        if fmt_tbls:
            self.data["loan_app_tables"] = fmt_tbls
            return 
        elif fmt_tbls == [] and extension in ['.doc', '.docx']:                    
            doc_path = file_path
            doc = Document(doc_path)
            # Iterate over all tables in the document
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tb_lst.append(table_data)
            fmt_tbls = self.format_tables(tb_lst)
            if fmt_tbls:
                self.data["loan_app_tables"] = fmt_tbls
                return
            else:
                if tb_lst:
                    self.data["loan_app_tables"] = tb_lst
                else:
                    result = converter.convert(file_path)
                    out = result.document.export_to_markdown()
                    fmt_tbls = out.split("Responses to the loan application")[0]
                    # fmt_tbls = re.findall(r'\|(.*?)\|', out)
                    self.data["loan_app_tables"] = fmt_tbls
                    return
        elif fmt_tbls == [] and extension in ['.pdf']:
            result = converter.convert(file_path)
            out = result.document.export_to_markdown()
            fmt_tbls = out.split("Responses to the loan application")[0]
            # fmt_tbls = re.findall(r'\|(.*?)\|', out)
            self.data["loan_app_tables"] = fmt_tbls
            return            

                            
    def img_to_txt(self, file_path: str, document_type: str) -> str:
        """
        Extracts text from an image or PDF file using Tesseract OCR.

        Args:
            file_path (str): Path to the image or PDF file.
            document_type (str): The type/category of the document.

        Returns:
            str: The extracted text from the file.
        """
        final_text = ""

        try:
            if file_path.lower().endswith('.pdf'):
                # Extract images from the PDF pages
                pages = convert_from_path(file_path)
                for page_number, page_image in enumerate(pages, start=1):
                    # Perform OCR on each page image
                    extracted_text_eng = pytesseract.image_to_string(page_image, lang='eng')
                    extracted_text_ara = pytesseract.image_to_string(page_image, lang='ara')
                    final_text += f"\n--- Page {page_number} ---\n{extracted_text_eng}\n{extracted_text_ara}\n"
            else:
                # Open the image file
                image = Image.open(file_path)

                # Perform OCR on the image
                extracted_text_eng = pytesseract.image_to_string(image, lang='eng')
                extracted_text_ara = pytesseract.image_to_string(image, lang='ara')

                final_text = f"{extracted_text_eng}\n{extracted_text_ara}"

            # Store the extracted text in the data dictionary
            self.data[document_type] = final_text
            return final_text

        except Exception as e:
            print(f"An error occurred while processing the file: {e}")
            return ""

    def market_data(self, file_path):
        # Load the Excel file
        df = pd.read_excel(file_path, sheet_name="Market Database  ", header=[0, 1])
        df = df.dropna(how='all')
        df = df.drop(index=20)
        df = df.ffill()
        # Initialize the dictionary
        market_data = []
        
        # Populate the dictionary
        for index, row in df.iterrows():
            data = {}
            data["Sector"] = row["Sector"].values[0]
            data["Product"] = row["Products "].values[0]
            data["Unit"] = row['Unit'].values[0]

            # Historical Sales Growth Rate
            data['Historical Sales Growth Rate'] = {
                '2020': f"{int(row['Historical Sales growth rate'][2020] * 100)}%",
                '2021': f"{int(row['Historical Sales growth rate'][2021] * 100)}%",
                '2022': f"{int(row['Historical Sales growth rate'][2022] * 100)}%",
                '2023': f"{int(row['Historical Sales growth rate'][2023] * 100)}%"
            }

            # Projected Sales Growth Rate
            data['Projected Sales Growth Rate'] = {
                '2024': f"{int(row['Projected Sales growth rate'][2024] * 100)}%",
                '2025': f"{int(row['Projected Sales growth rate'][2025] * 100)}%",
                '2026': f"{int(row['Projected Sales growth rate'][2026] * 100)}%",
                '2027': f"{int(row['Projected Sales growth rate'][2027] * 100)}%"
            }

            
            # Competitors Prices
            data['Competitors Prices'] = str(row["Competitors' prices (SR/unit)"].values[0])
            
            # Traders Importers Prices
            data['Traders Importers Prices'] = str(row["Traders/Importers prices (SR/unit)"].values[0])
            market_data.append(data)
        self.data["market_data"] = market_data

    def get(self, loan_app_name, cr_name, il_name):
        lapp_path = os.path.join(cwd, f'data/{loan_app_name}')
        il_path = os.path.join(cwd, f'data/{il_name}')
        cr_path = os.path.join(cwd, f'data/{cr_name}')
        md_path = os.path.join(cwd,'data/market_database.xlsx')

        self.loan_application_parser(file_path=lapp_path)
        print("D1")
        self.img_to_txt(file_path=il_path, document_type="industrial_licence")
        print("D2")
        self.img_to_txt(file_path=cr_path, document_type="commercial_registration")
        print("D3")
        self.market_data(file_path=md_path)
        print("D4")
        return self.data